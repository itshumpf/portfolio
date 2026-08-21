#!/usr/bin/env python3
"""
Rebuild Braeden Keena's DEPLOYED resume (braedenkeena.pages.dev/resume.pdf)
with ONLY the incorrect RF-fingerprinting figures corrected.

Source of truth for wording/structure/layout: the deployed PDF itself
(public/resume.pdf, LibreOffice 26.2 / Carlito, US Letter, 2pp).

Only change: the WiFi CSI bullet #2 figures.
  11-15 sigma            -> 12.7 sigma and 10.4 sigma, 6 of 6
  99.6% holdout accuracy -> 99.7% across 7,497 windows (single receiver),
                            95.7% across 14,234 windows (two receivers)
  2.8-million-frame      -> 2.36-million-frame

ASCII hyphens only. No non-breaking hyphens anywhere.
"""

import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION

# ---- tunables (points), matched against the deployed PDF ----
import os as _os
def _e(k,d): return float(_os.environ.get(k,d))
M_LEFT   = _e("M_LEFT", 45.00)
M_RIGHT  = _e("M_RIGHT", 44.40)
M_TOP    = _e("M_TOP", 35.00)
M_BOT    = 36.00
RIGHT_TAB = 450.20      # from left margin -> absolute x 495.3
BULLET_INDENT = 17.0
BULLET_HANG   = 10.0

FONT = "Carlito"

SZ_NAME = 20.0
SZ_SUB  = 11.0
SZ_CONTACT = 9.5
SZ_BODY = 10.5
SZ_SECTION = 11.5
SZ_JOB = 11.0
SZ_DATE = 10.0

SP_AFTER_BULLET = _e("SP_AB",2.0)
SP_BEFORE_JOB   = _e("SP_BJ",7.0)
SP_BEFORE_SECTION = _e("SP_BS",11.0)


def setfont(run, size, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    from docx.oxml.ns import qn
    rfonts.set(qn('w:ascii'), FONT)
    rfonts.set(qn('w:hAnsi'), FONT)
    rfonts.set(qn('w:cs'), FONT)
    rfonts.set(qn('w:eastAsia'), FONT)


def para(doc, space_before=0.0, space_after=0.0, align=None,
         left_indent=None, first_line=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.0
    pf.widow_control = False
    if align is not None:
        p.alignment = align
    if left_indent is not None:
        pf.left_indent = Pt(left_indent)
    if first_line is not None:
        pf.first_line_indent = Pt(first_line)
    return p


def bullet(doc, text, bold_lead=None, space_before=0.0,
           space_after=SP_AFTER_BULLET):
    p = para(doc, space_before=space_before, space_after=space_after,
             left_indent=BULLET_INDENT, first_line=-BULLET_HANG)
    p.paragraph_format.tab_stops.add_tab_stop(
        Pt(BULLET_INDENT), WD_TAB_ALIGNMENT.LEFT)
    setfont(p.add_run("•\t"), SZ_BODY)
    if bold_lead:
        setfont(p.add_run(bold_lead), SZ_BODY, bold=True)
    setfont(p.add_run(text), SZ_BODY)
    return p


def job(doc, title, date, space_before=SP_BEFORE_JOB):
    p = para(doc, space_before=space_before, space_after=0.0)
    p.paragraph_format.tab_stops.add_tab_stop(
        Pt(RIGHT_TAB), WD_TAB_ALIGNMENT.RIGHT)
    setfont(p.add_run(title), SZ_JOB, bold=True)
    setfont(p.add_run("\t"), SZ_JOB, bold=True)
    setfont(p.add_run(date), SZ_DATE, italic=True)
    return p


def section(doc, text, space_before=SP_BEFORE_SECTION):
    p = para(doc, space_before=space_before, space_after=0.0)
    setfont(p.add_run(text), SZ_SECTION, bold=True)
    return p


def build(path):
    doc = Document()
    s = doc.sections[0]
    s.page_width = Pt(612)
    s.page_height = Pt(792)
    s.left_margin = Pt(M_LEFT)
    s.right_margin = Pt(M_RIGHT)
    s.top_margin = Pt(M_TOP)
    s.bottom_margin = Pt(M_BOT)
    s.header_distance = Pt(0)
    s.footer_distance = Pt(0)

    st = doc.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(SZ_BODY)
    st.paragraph_format.space_after = Pt(0)
    st.paragraph_format.line_spacing = 1.0

    # ---------- header ----------
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    setfont(p.add_run("BRAEDEN KEENA"), SZ_NAME, bold=True)

    p = para(doc, space_before=_e("SP_SUB",2.0), align=WD_ALIGN_PARAGRAPH.CENTER)
    setfont(p.add_run(
        "Self-Taught Builder · Production Scrapers · "
        "Data Pipelines & Automation"), SZ_SUB, bold=True)

    p = para(doc, space_before=_e("SP_CON",1.5), align=WD_ALIGN_PARAGRAPH.CENTER)
    setfont(p.add_run(
        "Overland Park, KS 66213  ·  Braeden@thekeenas.com  ·  "
        "(303) 507-8626  ·  nacup.us  ·  findstorage.pages.dev"),
        SZ_CONTACT)

    # ---------- summary ----------
    p = para(doc, space_before=_e("SP_SUM",6.0))
    setfont(p.add_run(
        "Self-taught technical builder — I ship production scrapers, "
        "data pipelines, and automation. Three deployed, live web "
        "applications in the past 14 months, built end to end: API "
        "reverse-engineering, data pipelines, front-end, SEO, and daily "
        "automation. Backed by a decade of hands-on operations leadership "
        "managing teams of up to 100 people. I learn whatever a problem "
        "requires, then build the thing."), SZ_BODY)

    # ---------- shipped products ----------
    section(doc, "SHIPPED PRODUCTS (SOLO-BUILT, END TO END)")

    job(doc, "WiFi CSI RF-Fingerprinting Sensor Array — ESP32 + Python DSP",
        "2026 · Deployed / Live · Firmware / DSP", space_before=_e("SP_J1",9.7))
    bullet(doc,
        "Deployed and running 24/7 as the live, in-production home-security "
        "system for my own apartment — not a benchtop experiment. Built by "
        "applying the scientific method to a hard RF problem: hypothesized "
        "that a radio’s crystal-oscillator imperfections form a hardware "
        "fingerprint that survives MAC-address randomization, then proved "
        "through controlled experiments that device identity can be recovered "
        "from clock behavior rather than spoofable network identifiers.",
        space_before=_e("SP_B1",1.4))
    # ---- THE ONLY CORRECTED BULLET ----
    bullet(doc,
        "Built the full stack solo — ESP32 firmware (C / ESP-IDF) plus a "
        "Python DSP pipeline (RANSAC phase-slope fits → 1D Kalman drift "
        "tracking → Mahalanobis discrimination against per-source Gaussian "
        "models) — reaching 12.7σ and 10.4σ cross-manufacturer device "
        "separation on a 6/6 holdout, and 99.7% holdout accuracy across "
        "7,497 windows on a single receiver (95.7% across 14,234 windows "
        "spanning two receivers), over a 2.36-million-frame dataset.")
    bullet(doc,
        "Diagnosed thermal oscillator drift as the accuracy-limiting problem "
        "and engineered a novel reference-beacon drift-correction technique "
        "that cancels receiver-side drift and more than doubled device "
        "separation — turning an unstable measurement into a reliable "
        "fingerprint.")

    job(doc, "nacup.us — Live World Cup 2026 Fan Hub",
        "2026 · Founder / Developer")
    bullet(doc,
        "Designed, built, and deployed a full-stack production site in 48 "
        "hours; reached first-page organic search traffic within 24 hours of "
        "launch with 700+ visitors on day one during a live global event.",
        space_before=_e("SP_B1",1.4))
    bullet(doc,
        "Multi-source data pipeline (Google Places API, RSS, news scraping, "
        "live scores REST API) with dedupe/merge logic, trust-tier scoring, "
        "and freshness decay ranking across 38 host-city feeds.")
    bullet(doc,
        "Programmatically generated 80+ static SEO pages with JSON-LD "
        "structured data, sitemap, and Search Console integration; live-score "
        "polling with caching and fallback states.")
    bullet(doc,
        "Stack: Python, vanilla JavaScript, Leaflet.js, Cloudflare Pages, "
        "GitHub CI/CD.")

    job(doc, "findstorage.pages.dev — Nationwide Self-Storage Directory",
        "2025 · Founder / Developer")
    bullet(doc,
        "Daily self-scraping national self-storage pricing directory and "
        "market-research tool tracking ~4,600 stores / ~43,900 units — a "
        "7-pass discovery scraper with safety rails that aborts a publish if "
        "the store count drops below a floor or swings >10% run-over-run.",
        space_before=_e("SP_B1",1.4))
    bullet(doc,
        "Nationwide ZIP3 market analysis with matched-pairs attribute pricing "
        "(isolated a ~18–22% climate-control premium via "
        "same-store/same-size comparison), a “renter leverage” score, and "
        "per-store price-history drill-down.")
    bullet(doc,
        "Survived a real operator merger that added ~1,100 stores overnight "
        "without breaking; now on Cloudflare Pages, auto-refreshed daily via "
        "GitHub Actions.")

    job(doc, "Chrostory — Full-Stack Web App",
        "2025 · Founder (launched & sunset)")
    bullet(doc,
        "Took an original product from concept to public beta solo and "
        "validated it with 40+ real users; made the call to sunset it and "
        "applied the lessons — faster shipping, SEO-first architecture, "
        "automated data — directly to the two larger launches that followed.",
        space_before=_e("SP_B1",1.4))

    job(doc, "Discord Community Platform & Bots", "2019 – Present")
    bullet(doc,
        "Grew a community from zero to 3,500+ active members; built custom "
        "Python bots with NLP features to automate moderation and support "
        "e-commerce operations.", space_before=_e("SP_B1",1.4))
    bullet(doc,
        "Active in the KC Meshtastic (LoRa mesh networking) community — help "
        "plan node placement and RF coverage, troubleshoot member setups, and "
        "exchange design approaches with others building custom mesh networks.")

    # ---------- technical skills ----------
    section(doc, "TECHNICAL SKILLS")
    bullet(doc,
        "Python (scraping, data pipelines, automation), JavaScript, SQL, "
        "HTML/CSS", bold_lead="Languages: ", space_before=_e("SP_A3",6.7))
    bullet(doc,
        "REST APIs, API reverse-engineering, static site generation, "
        "technical SEO, Leaflet.js mapping, JSON data architecture",
        bold_lead="Web & Data: ")
    bullet(doc,
        "Git/GitHub, CI/CD (Cloudflare Pages, Netlify), Windows task "
        "automation, scheduled data refresh systems", bold_lead="DevOps: ")
    bullet(doc,
        "ESP32 firmware (C / ESP-IDF), ESP-NOW & BLE, I2C peripherals, "
        "sensor networks, 3D printing & CAD (OpenSCAD)",
        bold_lead="Hardware/IoT: ")

    # ---------- operations ----------
    section(doc, "OPERATIONS & LEADERSHIP EXPERIENCE")
    p = para(doc, space_before=_e("SP_A3",6.7))
    setfont(p.add_run(
        "10+ years of front-line operations and team leadership — including "
        "founding and operating my own business — managing teams of up to "
        "100 people across high-volume service environments, staffing, and "
        "vendor coordination."), SZ_BODY)

    job(doc, "Independent Technical Consultant", "7/2026 – Present")
    bullet(doc,
        "Build and ship production scrapers, data pipelines, and automation "
        "— turning messy public and web data sources into clean, scheduled, "
        "queryable datasets and dashboards (see shipped projects above).",
        space_before=_e("SP_B1",1.4))

    job(doc, "Crossing Guard Operations Supervisor — All City Management "
             "Services (ACMS)", "2023 – 2025")
    bullet(doc,
        "Hired, trained, and managed 50 crossing guards across 33 sites — 66 "
        "posts staffed twice every school day — holding ~25% annual turnover "
        "in a category known for severe churn.", space_before=_e("SP_B1",1.4))
    bullet(doc,
        "Coordinated daily across five constituencies (county sheriff, school "
        "principals, parents, the district, and the guards) and solved "
        "same-day coverage through a deep bench of on-call relationships — "
        "every student returned home safely, zero students injured across my "
        "full tenure.")

    job(doc, "Bartender — Tabard's Kitchen", "2019 – 2020")
    bullet(doc,
        "Ran the full lunch service solo as the only front-of-house staff — "
        "bartending, serving, taking and firing orders, running payments, and "
        "holding food-safety standards through the rush.", space_before=_e("SP_B1",1.4))

    job(doc, "Founder / Sole Operating Partner — Honey Lake Motocross Park, "
             "Milford, CA", "2016 – 2018")
    bullet(doc,
        "Acquired and revived a ~500-acre desert motocross park dormant since "
        "2007 — bought with two capital partners as the sole operating owner "
        "— and re-entitled it for racing after the prior owner had lost "
        "county approval and stripped the site: new survey, environmental "
        "review, a water well drilled under budget, and a heavy-equipment "
        "track rebuild.", space_before=_e("SP_B1",1.4))
    bullet(doc,
        "Directed a 120-person race-day crew built largely on in-kind pay "
        "with no payroll budget (peak ~6,500 attendees); designed and staffed "
        "the on-site medical program that executed a mid-event helicopter "
        "evacuation without incident.")
    bullet(doc,
        "Held full P&L — pricing, sponsorship, staffing, and marketing — and "
        "founded the Corey Herring Memorial race, which continued annually "
        "through 2023; exited by selling my stake.")

    # ---------- education ----------
    section(doc, "EDUCATION")
    p = para(doc, space_before=_e("SP_A3",6.7))
    setfont(p.add_run(
        "GED — State of Kansas, 2014. Everything since: self-taught, "
        "project-driven, and shipped to production."), SZ_BODY)

    doc.core_properties.author = "Braeden Keena"
    doc.core_properties.title = "Braeden Keena - Resume"
    doc.save(path)


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else "resume_source.docx"
    build(out)
    print("wrote", out)
